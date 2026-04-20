#!/usr/bin/env python3
"""Generate letterhead/quote DOCX, email signatures, and drawing PDFs for
each colour variant of the approved LK — Courier Slab logo.

Output lands in /assets/branding/final/:
    png/           — rasterized logo PNGs for DOCX embedding
    stationery/{variant}/letterhead.docx
    stationery/{variant}/quote.docx
    stationery/{variant}/email-signature.html
    drawings/drawing-{variant}.pdf

Variants: charcoal (standard), champagne (dark premium), mono (pure B&W),
two-tone (champagne + charcoal on warm mid-tone).
"""

import os
import cairosvg
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from PIL import Image, ImageDraw, ImageFont
import fitz  # PyMuPDF

COURIER_TTF = "/System/Library/Fonts/Supplemental/Courier New.ttf"
INTER_THIN_TTF = "/Users/dunderdoon/Library/Fonts/Inter-Thin.ttf"

ROOT = "/Users/dunderdoon/Projects_Local/primedesign"
FINAL = f"{ROOT}/assets/branding/final"
SVG_DIR = FINAL
PNG_DIR = f"{FINAL}/png"
STAT_DIR = f"{FINAL}/stationery"
DRAW_DIR = f"{FINAL}/drawings"
SOURCE_DRAW_PDF = f"{ROOT}/assets/branding/documents/drawing-options-prime-original.pdf"

GOLD = RGBColor(0xC9, 0xA9, 0x6E)
CHARCOAL = RGBColor(0x2C, 0x2C, 0x2C)
CHAMPAGNE = RGBColor(0xF2, 0xE6, 0xD0)
SLATE = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
NEAR_BLACK = RGBColor(0x0A, 0x0A, 0x0A)

CONTACT = {
    "name": "Lyda",
    "title": "Interior Designer | Project Manager",
    "phone": "0401 061 246",
    "email": "lyda@lkdesignandbuild.com.au",
    "website": "www.lkdesignandbuild.com.au",
    # TODO: swap to real PO Box string when Lyda has one.
    "address": "Adelaide, South Australia",
    "abn": "17 697 207 391",
}

TAGLINE = "Considered design. Precision build."

# One entry per colour variant.
VARIANTS = [
    {
        "key": "charcoal",
        "label": "Charcoal — Standard",
        "desc": "Charcoal LK on white. Standard business correspondence.",
        "logo_svg": "lk-logo-charcoal.svg",            # transparent
        "ink_hex": "#2C2C2C", "tag_hex": "#2C2C2C",
        "page_bg_hex": "FFFFFF",
        "page_bg_rgb": WHITE,
        "ink": CHARCOAL,
        "accent": GOLD,
        "body": CHARCOAL,
        "subtle": SLATE,
        "rule_hex": "C9A96E",
        "footer_is_dark": False,
        "email_bg": "#FFFFFF",
        "email_ink": "#2C2C2C",
        "email_accent": "#C9A96E",
    },
    {
        "key": "champagne",
        "label": "Champagne — Dark Premium",
        "desc": "Champagne LK on near-black. For proposal covers or dark-theme letterhead.",
        "logo_svg": "lk-logo-champagne.svg",
        "ink_hex": "#F2E6D0", "tag_hex": "#F2E6D0",
        "page_bg_hex": "0A0A0A",
        "page_bg_rgb": NEAR_BLACK,
        "ink": CHAMPAGNE,
        "accent": GOLD,
        "body": CHAMPAGNE,
        "subtle": RGBColor(0xB0, 0xA8, 0x9A),
        "rule_hex": "C9A96E",
        "footer_is_dark": True,
        "email_bg": "#0A0A0A",
        "email_ink": "#F2E6D0",
        "email_accent": "#C9A96E",
    },
    {
        "key": "mono",
        "label": "Mono — Black & White",
        "desc": "Pure black on white. Copier/fax safe, monochrome print.",
        "logo_svg": "lk-logo-mono-light.svg",          # has white BG rect
        "ink_hex": "#000000", "tag_hex": "#000000",
        "page_bg_hex": "FFFFFF",
        "page_bg_rgb": WHITE,
        "ink": RGBColor(0x00, 0x00, 0x00),
        "accent": RGBColor(0x00, 0x00, 0x00),
        "body": RGBColor(0x00, 0x00, 0x00),
        "subtle": RGBColor(0x44, 0x44, 0x44),
        "rule_hex": "000000",
        "footer_is_dark": False,
        "email_bg": "#FFFFFF",
        "email_ink": "#000000",
        "email_accent": "#000000",
    },
    {
        "key": "two-tone",
        "label": "Two-Tone — Warm Neutral",
        "desc": "Champagne LK + charcoal tagline on warm mid-tone. Branded marketing & stationery.",
        "logo_svg": "lk-logo-two-tone.svg",
        "ink_hex": "#F2E6D0", "tag_hex": "#2C2C2C",
        "page_bg_hex": "F5F0EB",                       # warm cream
        "page_bg_rgb": RGBColor(0xF5, 0xF0, 0xEB),
        "ink": CHARCOAL,
        "accent": GOLD,
        "body": CHARCOAL,
        "subtle": SLATE,
        "rule_hex": "C9A96E",
        "footer_is_dark": False,
        "email_bg": "#F5F0EB",
        "email_ink": "#2C2C2C",
        "email_accent": "#C9A96E",
    },
]


def rasterize_download_pngs():
    """Export a PNG alongside every downloadable SVG in /final so the hub can
    offer both formats from each logo tile. Written to /final/png/ with the
    same basename as the source SVG."""
    os.makedirs(PNG_DIR, exist_ok=True)
    # Every downloadable mark on the hub (Transparent / Solid / Favicon rows)
    # plus the gold horizontal used by the universal email signature.
    exports = [
        # Transparent
        "lk-logo-champagne.svg", "lk-logo-charcoal.svg", "lk-logo-two-tone.svg",
        "lk-logo-horizontal-champagne.svg", "lk-logo-horizontal-charcoal.svg", "lk-logo-horizontal-two-tone.svg",
        # Solid
        "lk-logo-dark.svg", "lk-logo-light.svg",
        "lk-logo-mono-dark.svg", "lk-logo-mono-light.svg",
        # Favicon
        "lk-icon-champagne.svg", "lk-icon-charcoal.svg", "lk-icon-two-tone.svg",
        # Signature-only (theme-agnostic accent gold)
        "lk-logo-horizontal-gold.svg",
    ]
    for name in exports:
        src = f"{SVG_DIR}/{name}"
        if not os.path.exists(src):
            print(f"  skipped (missing): {src}")
            continue
        out = f"{PNG_DIR}/{name[:-4]}.png"
        # 2400px wide is a sensible ceiling: crisp at any retina screen, still
        # < ~200 KB for text-based SVGs, and downscales cleanly to favicons.
        cairosvg.svg2png(url=src, write_to=out, output_width=2400)
        print(f"  rasterized {name} → {os.path.basename(out)}")


# Hex → 4-tuple helper for PIL
def _hex_to_rgba(hex_color, alpha=255):
    h = hex_color.lstrip('#')
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return (r, g, b, alpha)


def _draw_letter_spaced(draw, text, font, xy, fill, letter_spacing_px):
    """Draw text left-to-right applying letter-spacing after each glyph.
    xy is (x, y) anchor='ls' (left-baseline). Returns the ending x."""
    x, y = xy
    for ch in text:
        draw.text((x, y), ch, font=font, fill=fill, anchor="ls")
        x += font.getlength(ch) + letter_spacing_px
    return x


def render_letterhead_png(v, out_path):
    """Build the letterhead logo PNG directly in PIL so we can use the actual
    Inter Thin TTF for the tagline (cairosvg ignores font-weight and has no
    reliable path to select a specific face on macOS).

    viewBox "210 210 280 194" → canvas 3000×2079 px (scale F ≈ 10.71 px/vb).
    LK uses Courier New at fs=240 (vb), letter-spacing −0.02em, baseline y=350.
    Tagline uses Inter Thin at fs=28 (vb), letter-spacing 0.1, baseline y=400.
    """
    W, H = 3000, 2079
    F = W / 280.0  # px per viewBox unit

    img = Image.new('RGBA', (W, H), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)

    # LK
    lk_font = ImageFont.truetype(COURIER_TTF, int(round(240 * F)))
    ink_rgba = _hex_to_rgba(v["ink_hex"])
    lk_x = (200 - 210) * F  # text anchor x, in PNG coords
    lk_y = (350 - 210) * F  # baseline
    ls_lk = -0.02 * 240 * F  # -4.8 vb → px
    _draw_letter_spaced(draw, "LK", lk_font, (lk_x, lk_y), ink_rgba, ls_lk)

    # Tagline — Inter Thin. Calibrated against the actual Courier LK render:
    # measured LK visible left = 214.95 vb, right = 477.03 vb (width 262.08).
    # Inter-Thin D lsb = 3.17 vb, rsb = 2.15 vb, advance sum = 258.6 vb.
    #   draw x = 214.95 − 3.17 = 211.78
    #   letter-spacing = (262.08 − (258.6 − 3.17 − 2.15)) / 15 = 0.587 vb/gap
    tag_font = ImageFont.truetype(INTER_THIN_TTF, int(round(28 * F)))
    tag_rgba = _hex_to_rgba(v["tag_hex"], int(255 * 0.55))  # fainter opacity
    tg_x = (211.78 - 210) * F
    tg_y = (400 - 210) * F
    ls_tg = 0.587 * F
    _draw_letter_spaced(draw, "DESIGN AND BUILD", tag_font, (tg_x, tg_y), tag_rgba, ls_tg)

    img.save(out_path, 'PNG', optimize=True)


def rasterize_logos():
    """Build each variant's letterhead PNG via PIL (Inter Thin) and the
    horizontal nav PNG via cairosvg."""
    os.makedirs(PNG_DIR, exist_ok=True)
    for v in VARIANTS:
        # Letterhead: PIL rendering with real Inter Thin.
        letterhead_png = f"{PNG_DIR}/{v['key']}-letterhead.png"
        render_letterhead_png(v, letterhead_png)
        v["logo_letterhead_png"] = letterhead_png
        print(f"  PIL-rendered letterhead → {letterhead_png}")

        # Horizontal variant (email signature / nav) — cairosvg is fine, only
        # LK is Courier New which renders correctly, and the horizontal tagline
        # is tiny enough that the weight doesn't matter visually.
        horiz_svg = f"{SVG_DIR}/{v['logo_svg']}".replace("lk-logo-", "lk-logo-horizontal-")
        if not os.path.exists(horiz_svg):
            horiz_svg = f"{SVG_DIR}/{v['logo_svg']}"
        horiz_png = f"{PNG_DIR}/{v['key']}.png"
        cairosvg.svg2png(url=horiz_svg, write_to=horiz_png, output_width=1800)
        v["logo_png"] = horiz_png
        print(f"  rasterized {os.path.basename(horiz_svg)} → {horiz_png}")


def set_page_shading(doc, hex_color):
    """Apply a page background colour to an entire section."""
    if hex_color.upper() == "FFFFFF":
        return
    sect = doc.sections[0]
    sectPr = sect._sectPr
    # background fill via <w:background>
    bg = parse_xml(f'<w:background {nsdecls("w")} w:color="{hex_color}"/>')
    doc.element.insert(0, bg)
    # also tell Word to display the bg
    settings = doc.settings.element
    dispBg = parse_xml(f'<w:displayBackgroundShape {nsdecls("w")}/>')
    settings.append(dispBg)


def style_run(run, font_name, size, color, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def rule(doc, hex_color, tight=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0 if tight else 2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="{hex_color}"/></w:pBdr>'
    )
    pPr.append(pBdr)


def add_logo_header(doc, v):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    # Letterhead variant: LK + readable tagline beneath, left-aligned.
    run.add_picture(v["logo_letterhead_png"], height=Cm(2.6))
    rule(doc, v["rule_hex"], tight=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    contact = f'{CONTACT["phone"]}  |  {CONTACT["email"]}  |  {CONTACT["website"]}  |  {CONTACT["address"]}'
    run = p.add_run(contact)
    style_run(run, "Calibri", 8, v["subtle"])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_footer(doc, v):
    rule(doc, v["rule_hex"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f'LK Design and Build  |  ABN {CONTACT["abn"]}  |  {TAGLINE}')
    style_run(run, "Calibri", 7.5, v["subtle"], italic=True)


def build_letterhead(v):
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2); s.bottom_margin = Cm(2); s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
    set_page_shading(doc, v["page_bg_hex"])
    add_logo_header(doc, v)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(12)
    style_run(p.add_run("[Date]"), "Calibri", 10, v["body"])
    for line in ["[Recipient Name]", "[Company / Address Line 1]", "[Address Line 2]", "[City, State, Postcode]"]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        style_run(p.add_run(line), "Calibri", 10, v["body"])
    doc.add_paragraph()
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    style_run(p.add_run("Re: [Subject Line]"), "Calibri", 10, v["body"], bold=True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    style_run(p.add_run("Dear [Name],"), "Calibri", 10, v["body"])

    body_text = (
        "Thank you for your interest in LK Design and Build. We specialise in premium home renovations, "
        "additions, and new builds across Adelaide, delivering exceptional craftsmanship and thoughtful "
        "design at every stage.\n\n"
        "[Continue your letter here. This template provides the branded letterhead with the approved LK "
        "— Courier Slab logo. Simply replace the placeholder text with your content.]\n\n"
        "We look forward to discussing your project in detail."
    )
    for para in body_text.split("\n\n"):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8); p.paragraph_format.line_spacing = Pt(14)
        style_run(p.add_run(para), "Calibri", 10, v["body"])

    doc.add_paragraph()
    p = doc.add_paragraph(); style_run(p.add_run("Kind regards,"), "Calibri", 10, v["body"])
    doc.add_paragraph(); doc.add_paragraph()
    p = doc.add_paragraph(); style_run(p.add_run("Lyda"), "Calibri", 10, v["body"], bold=True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run("LK Design and Build"), "Calibri", 9, v["accent"])

    add_footer(doc, v)
    out = f"{STAT_DIR}/{v['key']}/letterhead.docx"
    doc.save(out)
    print(f"  letterhead: {out}")


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def build_quote(v):
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2); s.bottom_margin = Cm(2); s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
    set_page_shading(doc, v["page_bg_hex"])
    add_logo_header(doc, v)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run("QUOTATION"), "Calibri", 20, v["accent"], bold=True)
    rule(doc, v["rule_hex"])

    details = [
        ("Quote Number:", "[QTE-0001]"),
        ("Date:", "[Date]"),
        ("Valid Until:", "[Date + 30 days]"),
        ("Project Address:", "[Site Address]"),
    ]
    t = doc.add_table(rows=len(details), cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (lbl, val) in enumerate(details):
        c1, c2 = t.cell(i, 0), t.cell(i, 1)
        p = c1.paragraphs[0]; style_run(p.add_run(lbl), "Calibri", 9, v["subtle"], bold=True); p.paragraph_format.space_after = Pt(3)
        p = c2.paragraphs[0]; style_run(p.add_run(val), "Calibri", 9, v["body"]); p.paragraph_format.space_after = Pt(3)
        c1.width = Cm(4); c2.width = Cm(12)
    _no_borders(t)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run("PREPARED FOR"), "Calibri", 9, v["accent"])
    for line in ["[Client Name]", "[Client Address]", "[Client Email]", "[Client Phone]"]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
        style_run(p.add_run(line), "Calibri", 9, v["body"])

    doc.add_paragraph()
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run("SCOPE OF WORKS"), "Calibri", 9, v["accent"])
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    style_run(p.add_run("[Describe proposed works, demolition, structural, finishes, fixtures, allowances.]"),
              "Calibri", 9, v["subtle"], italic=True)

    items_table = doc.add_table(rows=1, cols=3); items_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Description", "Qty", "Amount (ex GST)"]
    widths = [Cm(10), Cm(2), Cm(4.5)]
    header_bg = "2C2C2C" if not v["footer_is_dark"] else "F2E6D0"
    header_fg = WHITE if not v["footer_is_dark"] else NEAR_BLACK
    for i, (h, w) in enumerate(zip(headers, widths)):
        c = items_table.cell(0, i); c.width = w
        p = c.paragraphs[0]; style_run(p.add_run(h), "Calibri", 8.5, header_fg, bold=True)
        set_cell_shading(c, header_bg)
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)

    samples = [
        ("[Demolition & Site Preparation]", "1", "$X,XXX.XX"),
        ("[Structural Works]", "1", "$X,XXX.XX"),
        ("[Carpentry & Framing]", "1", "$X,XXX.XX"),
        ("[Electrical]", "1", "$X,XXX.XX"),
        ("[Plumbing]", "1", "$X,XXX.XX"),
        ("[Tiling & Waterproofing]", "1", "$X,XXX.XX"),
        ("[Joinery & Cabinetry]", "1", "$X,XXX.XX"),
        ("[Painting & Finishes]", "1", "$X,XXX.XX"),
        ("[Fixtures & Fittings — Allowance]", "1", "$X,XXX.XX"),
        ("[Project Management & Supervision]", "1", "$X,XXX.XX"),
    ]
    for desc, qty, amt in samples:
        row = items_table.add_row()
        for i, val in enumerate([desc, qty, amt]):
            c = row.cells[i]; c.width = widths[i]
            p = c.paragraphs[0]; style_run(p.add_run(val), "Calibri", 9, v["body"])
            p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
            if i == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _items_borders(items_table, v["rule_hex"])

    doc.add_paragraph()
    totals = [("Subtotal (ex GST):", "$XX,XXX.XX"),
              ("GST (10%):", "$X,XXX.XX"),
              ("TOTAL (inc GST):", "$XX,XXX.XX")]
    for lbl, val in totals:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_after = Pt(2)
        is_total = "TOTAL" in lbl and "Sub" not in lbl
        style_run(p.add_run(lbl + "  "), "Calibri", 10 if is_total else 9, v["body"], bold=is_total)
        style_run(p.add_run(val), "Calibri", 10 if is_total else 9, v["accent"] if is_total else v["body"], bold=is_total)

    rule(doc, v["rule_hex"]); doc.add_paragraph()
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run("TERMS & CONDITIONS"), "Calibri", 9, v["accent"])
    terms = [
        "This quotation is valid for 30 days from the date of issue.",
        "A 10% deposit is required upon acceptance to secure scheduling.",
        "Progress payments are due at agreed milestones as outlined in the building contract.",
        "All works are carried out in accordance with the Building Code of Australia and relevant South Australian regulations.",
        "Any variations to the scope of works will be documented and agreed in writing before commencement.",
        "LK Design and Build holds comprehensive public liability and construction insurance.",
        "Defects liability period: 12 months from practical completion.",
    ]
    for i, term in enumerate(terms, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.left_indent = Cm(0.5)
        style_run(p.add_run(f"{i}. {term}"), "Calibri", 8, v["subtle"])

    doc.add_paragraph()
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run("ACCEPTANCE"), "Calibri", 9, v["accent"])
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run("I/We accept this quotation and the terms and conditions outlined above."), "Calibri", 9, v["body"])

    doc.add_paragraph()
    sig = doc.add_table(rows=2, cols=2); sig.alignment = WD_TABLE_ALIGNMENT.LEFT
    labels = [("Client Signature:", "Date:"), ("Print Name:", "")]
    for r, (l1, l2) in enumerate(labels):
        for c, lbl in enumerate([l1, l2]):
            if not lbl: continue
            cell = sig.cell(r, c)
            p = cell.paragraphs[0]
            style_run(p.add_run(lbl + "  ___________________________"), "Calibri", 9, v["body"])
            p.paragraph_format.space_before = Pt(12)
    _no_borders(sig)

    add_footer(doc, v)
    out = f"{STAT_DIR}/{v['key']}/quote.docx"
    doc.save(out)
    print(f"  quote:      {out}")


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _items_borders(table, hex_color):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'<w:insideH w:val="single" w:sz="2" w:space="0" w:color="E8E8E8"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


EMAIL_TEMPLATE = """<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>Email Signature — LK Design and Build</title></head>
<body style="margin:0;padding:20px;background:#f5f5f5;font-family:Arial,sans-serif;">
<div style="max-width:620px;margin:0 auto;background:#fff;padding:30px;border-radius:8px;">
<div style="display:flex;align-items:center;justify-content:space-between;gap:12px;margin-bottom:12px;flex-wrap:wrap;">
  <p style="font-size:12px;color:#666;margin:0;flex:1;min-width:260px;line-height:1.55;">Click <strong>Copy signature</strong>, then paste into Gmail (Settings → Signature) or Outlook (File → Options → Mail → Signatures). Copies as rich HTML — logo, link colours and layout are preserved.</p>
  <button id="copy-btn" type="button" onclick="copySignature()" style="font-family:Inter,Arial,sans-serif;font-size:12px;font-weight:600;letter-spacing:0.04em;padding:9px 16px;background:#C9A96E;color:#fff;border:none;border-radius:6px;cursor:pointer;display:inline-flex;align-items:center;gap:6px;transition:background 0.2s;">
    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round"><rect x="9" y="9" width="13" height="13" rx="2" ry="2"/><path d="M5 15H4a2 2 0 01-2-2V4a2 2 0 012-2h9a2 2 0 012 2v1"/></svg>
    <span id="copy-btn-label">Copy signature</span>
  </button>
</div>

<div style="font-size:11px;color:#888;background:#fafaf7;border:1px solid #eee3d0;border-radius:6px;padding:10px 14px;margin-bottom:18px;line-height:1.55;">
  <strong style="color:#6b5a3a;">Theme-agnostic design.</strong> No background colour is copied, and the mark + text use gold (<code>#C9A96E</code>) and medium charcoal (<code>#4A4A4A</code>) — both read on light-mode and dark-mode email clients. Preview below is on white to match the default Gmail / Outlook compose view.
</div>

<!-- Sample swatch so Lyda can see how it looks on both themes -->
<div style="display:grid;grid-template-columns:1fr 1fr;gap:0;border-radius:6px;overflow:hidden;border:1px solid #eee;margin-bottom:6px;">
  <div style="background:#FFFFFF;padding:10px 14px 6px;">
    <div style="font-size:9px;color:#999;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:6px;">Light mode preview</div>
{sig_light}
  </div>
  <div style="background:#1B1B1B;padding:10px 14px 6px;">
    <div style="font-size:9px;color:#888;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:6px;">Dark mode preview</div>
{sig_dark}
  </div>
</div>
<p style="font-size:10px;color:#aaa;margin:4px 0 18px;text-align:center;">Identical signature rendered on both backgrounds — the Copy button below copies this exact markup (no background).</p>

<hr style="border:1px dashed #ccc;margin:0 0 12px;">

<!-- ====== EMAIL SIGNATURE START — this is what gets copied ====== -->
<div id="signature">
<table cellpadding="0" cellspacing="0" border="0" style="font-family:Calibri,Arial,sans-serif;color:#6A6A6A;">
<tr>
  <td style="vertical-align:top;padding:0 18px 0 0;border-right:1px solid #C9A96E;">
    <img src="https://dunderdoon.com/assets/branding/final/png/lk-logo-horizontal-gold.png" alt="LK Design and Build" style="width:160px;height:auto;display:block;">
  </td>
  <td style="vertical-align:top;padding:0 0 0 18px;">
    <div style="font-size:15px;font-weight:bold;color:#6A6A6A;margin-bottom:1px;">Lyda</div>
    <div style="font-size:10px;color:#C9A96E;font-weight:600;letter-spacing:1.2px;margin-bottom:10px;">{title}</div>
    <div style="font-size:11px;color:#6A6A6A;line-height:1.7;">
      <span>M:</span> {phone}<br>
      <span>E:</span> <a href="mailto:{email}" style="color:#C9A96E;text-decoration:none;">{email}</a><br>
      <span>A:</span> {address}
    </div>
    <div style="margin-top:8px;padding-top:8px;border-top:1px solid #C9A96E;font-size:9px;color:#9A9A9A;font-style:italic;">
      {tagline}<br><span style="font-style:normal;letter-spacing:0.5px;color:#9A9A9A;">ABN {abn}</span>
    </div>
  </td>
</tr>
</table>
</div>
<!-- ====== EMAIL SIGNATURE END ====== -->
</div>

<script>
async function copySignature() {{
  var sig = document.getElementById('signature');
  var btn = document.getElementById('copy-btn');
  var label = document.getElementById('copy-btn-label');
  var original = label.textContent;

  var html = sig.innerHTML;
  var text = sig.innerText;
  var ok = false;

  try {{
    if (navigator.clipboard && window.ClipboardItem) {{
      var item = new ClipboardItem({{
        'text/html': new Blob([html], {{type: 'text/html'}}),
        'text/plain': new Blob([text], {{type: 'text/plain'}})
      }});
      await navigator.clipboard.write([item]);
      ok = true;
    }}
  }} catch (e) {{ /* fall through */ }}

  if (!ok) {{
    try {{
      var range = document.createRange();
      range.selectNodeContents(sig);
      var sel = window.getSelection();
      sel.removeAllRanges();
      sel.addRange(range);
      ok = document.execCommand('copy');
      sel.removeAllRanges();
    }} catch (e) {{ ok = false; }}
  }}

  label.textContent = ok ? 'Copied ✓' : 'Copy failed — select manually';
  btn.style.background = ok ? '#7FA66D' : '#C14A4A';
  setTimeout(function() {{
    label.textContent = original;
    btn.style.background = '#C9A96E';
  }}, 2200);
}}
</script>
</body></html>
"""

SIG_SAMPLE = """<table cellpadding="0" cellspacing="0" border="0" style="font-family:Calibri,Arial,sans-serif;color:#6A6A6A;">
<tr>
  <td style="vertical-align:top;padding:0 14px 0 0;border-right:1px solid #C9A96E;">
    <img src="https://dunderdoon.com/assets/branding/final/png/lk-logo-horizontal-gold.png" alt="" style="width:130px;height:auto;display:block;">
  </td>
  <td style="vertical-align:top;padding:0 0 0 14px;">
    <div style="font-size:13px;font-weight:bold;color:#6A6A6A;margin-bottom:1px;">Lyda</div>
    <div style="font-size:9px;color:#C9A96E;font-weight:600;letter-spacing:1.2px;margin-bottom:8px;">{title}</div>
    <div style="font-size:10px;color:#6A6A6A;line-height:1.6;">
      <span>M:</span> {phone}<br>
      <span>E:</span> <a href="mailto:{email}" style="color:#C9A96E;text-decoration:none;">{email}</a><br>
      <span>A:</span> {address}
    </div>
    <div style="margin-top:6px;padding-top:6px;border-top:1px solid #C9A96E;font-size:8px;color:#9A9A9A;font-style:italic;">
      {tagline}<br><span style="font-style:normal;letter-spacing:0.4px;color:#9A9A9A;">ABN {abn}</span>
    </div>
  </td>
</tr>
</table>"""


def build_email_signature(v):
    # Theme-agnostic signature — identical output regardless of variant.
    # Both previews use the same colours as what actually gets copied, so
    # Lyda sees the honest result on both light and dark backgrounds.
    fields = dict(
        title=CONTACT["title"], phone=CONTACT["phone"], email=CONTACT["email"],
        address=CONTACT["address"], abn=CONTACT["abn"], tagline=TAGLINE,
    )
    sig_light = SIG_SAMPLE.format(**fields)
    sig_dark  = SIG_SAMPLE.format(**fields)
    html = EMAIL_TEMPLATE.format(sig_light=sig_light, sig_dark=sig_dark, **fields)
    out = f"{STAT_DIR}/{v['key']}/email-signature.html"
    with open(out, "w") as f:
        f.write(html)
    print(f"  email:      {out}")


def build_drawing(v):
    """Overlay the LK logo in place of the Prime title-block logo."""
    if not os.path.exists(SOURCE_DRAW_PDF):
        print(f"  drawing skipped: source PDF missing at {SOURCE_DRAW_PDF}")
        return
    doc = fitz.open(SOURCE_DRAW_PDF)
    page = doc[0]
    # Locate the Prime logo image — heuristic: largest raster in the title block (bottom-right).
    images = page.get_images(full=True)
    if not images:
        print("  drawing skipped: no images on page 1")
        doc.close(); return
    # Pick the image whose bbox is in the bottom-right quadrant of the page.
    page_rect = page.rect
    target_xref = None
    target_rect = None
    for img in images:
        xref = img[0]
        for rect in page.get_image_rects(xref):
            if rect.x1 > page_rect.width * 0.55 and rect.y0 > page_rect.height * 0.55:
                if target_rect is None or (rect.width * rect.height) > (target_rect.width * target_rect.height):
                    target_xref = xref; target_rect = rect
    if target_rect is None:
        # Fallback: pick the largest overall
        for img in images:
            xref = img[0]
            for rect in page.get_image_rects(xref):
                if target_rect is None or (rect.width * rect.height) > (target_rect.width * target_rect.height):
                    target_xref = xref; target_rect = rect
    # Hide original with a solid rect matching page bg, then draw new logo on top.
    bg_rgb = tuple(c / 255 for c in v["page_bg_rgb"])
    page.draw_rect(target_rect, color=bg_rgb, fill=bg_rgb, width=0)
    # Insert new logo centred in the same rect, preserving aspect.
    page.insert_image(target_rect, filename=v["logo_png"], keep_proportion=True)
    out = f"{DRAW_DIR}/drawing-{v['key']}.pdf"
    doc.save(out)
    doc.close()
    print(f"  drawing:    {out}")


def main():
    rasterize_download_pngs()
    rasterize_logos()
    for v in VARIANTS:
        os.makedirs(f"{STAT_DIR}/{v['key']}", exist_ok=True)
        print(f"\n=== {v['label']} ===")
        build_letterhead(v)
        build_quote(v)
        build_email_signature(v)
        build_drawing(v)
    print(f"\nDone. Generated stationery for {len(VARIANTS)} colour variants.")


if __name__ == "__main__":
    main()
